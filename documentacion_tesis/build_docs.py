"""Convertir los MDs de documentacion a HTML y DOCX.

Uso:
    uv run --index-url https://pypi.ci.artifacts.walmart.com/artifactory/api/pypi/external-pypi/simple \\
        --allow-insecure-host pypi.ci.artifacts.walmart.com \\
        --with markdown --with python-docx \\
        python build_docs.py
"""

from __future__ import annotations

import re
from pathlib import Path

import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm

BASE_DIR = Path(__file__).parent
FILES = [
    ("documentacion_narrativa.md", "Documentación Narrativa Académica"),
    ("documentacion_tecnica.md", "Documentación Técnica"),
    ("documentacion_simple.md", "Documentación para Todo Público"),
]


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<meta name="color-scheme" content="light dark">
<title>{title}</title>
<style>
:root {{
  --primary: #0E7C7B;
  --primary-dark: #095958;
  --accent: #F26430;
  --bg: #FFFFFF;
  --bg-soft: #F7F5F0;
  --text: #1A1A1A;
  --text-muted: #5A5A5A;
  --border: #E1DED6;
  --code-bg: #F2F0EA;
  --shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
}}

@media (prefers-color-scheme: dark) {{
  :root:not([data-theme="light"]) {{
    --primary: #4DBAB9;
    --primary-dark: #7FD4D3;
    --accent: #FF8B5C;
    --bg: #1A1A1A;
    --bg-soft: #242424;
    --text: #E8E6E1;
    --text-muted: #A8A5A0;
    --border: #3A3A3A;
    --code-bg: #2A2A2A;
    --shadow: 0 2px 8px rgba(0, 0, 0, 0.4);
  }}
}}

[data-theme="dark"] {{
  --primary: #4DBAB9;
  --primary-dark: #7FD4D3;
  --accent: #FF8B5C;
  --bg: #1A1A1A;
  --bg-soft: #242424;
  --text: #E8E6E1;
  --text-muted: #A8A5A0;
  --border: #3A3A3A;
  --code-bg: #2A2A2A;
  --shadow: 0 2px 8px rgba(0, 0, 0, 0.4);
}}

* {{ box-sizing: border-box; }}
html, body {{ margin: 0; padding: 0; }}
body {{
  font-family: Georgia, "Times New Roman", serif;
  font-size: 17px;
  line-height: 1.65;
  color: var(--text);
  background: var(--bg);
  max-width: 820px;
  margin: 0 auto;
  padding: 48px 32px 96px;
  transition: background 0.2s ease, color 0.2s ease;
}}

.theme-toggle {{
  position: fixed;
  top: 16px;
  right: 16px;
  z-index: 100;
  width: 44px;
  height: 44px;
  border-radius: 50%;
  border: 1px solid var(--border);
  background: var(--bg-soft);
  color: var(--text);
  cursor: pointer;
  font-size: 18px;
  display: flex;
  align-items: center;
  justify-content: center;
  box-shadow: var(--shadow);
  transition: transform 0.15s ease, background 0.2s ease, border-color 0.2s ease;
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
}}
.theme-toggle:hover {{
  transform: scale(1.08);
  border-color: var(--primary);
}}
.theme-toggle:focus-visible {{
  outline: 2px solid var(--primary);
  outline-offset: 2px;
}}
.theme-toggle .icon-sun,
.theme-toggle .icon-moon {{ display: block; }}
.theme-toggle .icon-sun {{ display: none; }}
[data-theme="dark"] .theme-toggle .icon-sun {{ display: block; }}
[data-theme="dark"] .theme-toggle .icon-moon {{ display: none; }}
@media (prefers-color-scheme: dark) {{
  :root:not([data-theme="light"]) .theme-toggle .icon-sun {{ display: block; }}
  :root:not([data-theme="light"]) .theme-toggle .icon-moon {{ display: none; }}
}}
@media print {{
  .theme-toggle {{ display: none; }}
}}
h1, h2, h3, h4 {{
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
  color: var(--primary-dark);
  line-height: 1.25;
  margin-top: 2em;
  margin-bottom: 0.6em;
}}
h1 {{
  font-size: 2.4em;
  border-bottom: 3px solid var(--primary);
  padding-bottom: 12px;
  margin-top: 0;
}}
h2 {{
  font-size: 1.7em;
  border-bottom: 1px solid var(--border);
  padding-bottom: 6px;
}}
h3 {{ font-size: 1.3em; color: var(--primary); }}
h4 {{ font-size: 1.1em; color: var(--text); }}
p {{ margin: 0.8em 0; }}
strong {{ color: var(--primary-dark); }}
a {{ color: var(--primary); text-decoration: underline; }}
a:hover {{ color: var(--accent); }}
ul, ol {{ padding-left: 1.6em; }}
li {{ margin: 0.25em 0; }}
code {{
  font-family: "SF Mono", Consolas, "Courier New", monospace;
  font-size: 0.9em;
  background: var(--code-bg);
  padding: 2px 6px;
  border-radius: 3px;
  color: var(--primary-dark);
  border: 1px solid var(--border);
}}
pre {{
  background: var(--code-bg);
  border-left: 4px solid var(--primary);
  padding: 16px 20px;
  overflow-x: auto;
  border-radius: 4px;
  font-size: 0.85em;
  line-height: 1.5;
}}
pre code {{
  background: transparent;
  padding: 0;
  color: var(--text);
  border: none;
}}
blockquote {{
  border-left: 4px solid var(--accent);
  padding-left: 16px;
  margin: 1em 0;
  color: var(--text-muted);
  font-style: italic;
}}
table {{
  border-collapse: collapse;
  margin: 1em 0;
  width: 100%;
  font-size: 0.92em;
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
}}
th, td {{
  border: 1px solid var(--border);
  padding: 8px 12px;
  text-align: left;
  vertical-align: top;
}}
th {{
  background: var(--bg-soft);
  color: var(--primary-dark);
  font-weight: 600;
}}
tbody tr:nth-child(even) {{
  background: var(--bg-soft);
}}
hr {{
  border: none;
  border-top: 1px solid var(--border);
  margin: 2.5em 0;
}}
.doc-meta {{
  color: var(--text-muted);
  font-size: 0.9em;
  font-style: italic;
  text-align: center;
  margin-bottom: 32px;
}}
@media print {{
  body {{ font-size: 11pt; padding: 0; }}
  h1 {{ page-break-before: always; }}
  h1:first-of-type {{ page-break-before: avoid; }}
  h2, h3 {{ page-break-after: avoid; }}
  pre, table, blockquote {{ page-break-inside: avoid; }}
}}
@media (max-width: 640px) {{
  body {{ padding: 24px 16px; font-size: 16px; }}
  h1 {{ font-size: 1.9em; }}
  h2 {{ font-size: 1.4em; }}
}}
</style>
</head>
<body>
<button class="theme-toggle" id="themeToggle" type="button" aria-label="Alternar tema claro/oscuro" title="Alternar tema claro/oscuro">
  <svg class="icon-moon" xmlns="http://www.w3.org/2000/svg" width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M21 12.79A9 9 0 1 1 11.21 3 7 7 0 0 0 21 12.79z"></path></svg>
  <svg class="icon-sun" xmlns="http://www.w3.org/2000/svg" width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><circle cx="12" cy="12" r="4"></circle><path d="M12 2v2M12 20v2M4.93 4.93l1.41 1.41M17.66 17.66l1.41 1.41M2 12h2M20 12h2M4.93 19.07l1.41-1.41M17.66 6.34l1.41-1.41"></path></svg>
</button>
{body}
<script>
(function () {{
  var STORAGE_KEY = 'servel-docs-theme';
  var root = document.documentElement;
  var btn = document.getElementById('themeToggle');
  var stored = null;
  try {{ stored = localStorage.getItem(STORAGE_KEY); }} catch (e) {{}}
  if (stored === 'light' || stored === 'dark') {{
    root.setAttribute('data-theme', stored);
  }}
  function currentEffectiveTheme() {{
    var attr = root.getAttribute('data-theme');
    if (attr === 'light' || attr === 'dark') return attr;
    return window.matchMedia('(prefers-color-scheme: dark)').matches ? 'dark' : 'light';
  }}
  btn.addEventListener('click', function () {{
    var next = currentEffectiveTheme() === 'dark' ? 'light' : 'dark';
    root.setAttribute('data-theme', next);
    try {{ localStorage.setItem(STORAGE_KEY, next); }} catch (e) {{}}
  }});
}})();
</script>
</body>
</html>
"""


def md_to_html(md_path: Path, out_path: Path, title: str) -> None:
    text = md_path.read_text(encoding="utf-8")
    html_body = markdown.markdown(
        text,
        extensions=["extra", "toc", "sane_lists", "tables", "fenced_code"],
    )
    html = HTML_TEMPLATE.format(title=title, body=html_body)
    out_path.write_text(html, encoding="utf-8")
    print(f"HTML generado: {out_path.name}")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

HEADING_COLOR = RGBColor(0x09, 0x59, 0x58)
BODY_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
CODE_COLOR = RGBColor(0x33, 0x33, 0x33)
MUTED_COLOR = RGBColor(0x5A, 0x5A, 0x5A)

# Academic formatting defaults (APA-style)
BODY_FONT = "Times New Roman"
BODY_FONT_SIZE = 12  # pt
MARGIN_CM = 3.0      # 3 cm on all sides


def _insert_toc_field(doc: Document) -> None:
    """Insert a native Word TOC field. User can right-click > Update Field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Tabla de contenidos: click derecho > Actualizar campos para poblar."
    )
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _add_page_number_footer(section, doc_title: str) -> None:
    """Center 'Titulo | Pagina N' in the footer of a section."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = p.add_run(f"{doc_title}  |  Página ")
    prefix.font.name = BODY_FONT
    prefix.font.size = Pt(9)
    prefix.font.italic = True
    prefix.font.color.rgb = MUTED_COLOR
    run = p.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED_COLOR
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _configure_document(doc: Document, doc_title: str) -> None:
    """Apply academic formatting: margins, Normal style, header/footer."""
    for section in doc.sections:
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)
        _add_page_number_footer(section, doc_title)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_FONT_SIZE)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    # Ensure Heading styles inherit Times New Roman (Word usually applies Calibri).
    for level in (1, 2, 3, 4):
        try:
            style = doc.styles[f"Heading {level}"]
            style.font.name = BODY_FONT
            style.font.color.rgb = HEADING_COLOR
            style.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except KeyError:
            pass


def _set_font(run, name: str = BODY_FONT, size: int = BODY_FONT_SIZE,
              bold: bool = False, italic: bool = False,
              color: RGBColor = BODY_COLOR) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int) -> None:
    """Add heading using native Word 'Heading N' styles so TOC field detects it."""
    sizes = {1: 18, 2: 14, 3: 12, 4: 12}
    safe_level = min(max(level, 1), 4)
    p = doc.add_paragraph(style=f"Heading {safe_level}")
    p.paragraph_format.space_before = Pt(18 if safe_level <= 2 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    _set_font(run, name=BODY_FONT, size=sizes[safe_level], bold=True,
              color=HEADING_COLOR)


INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _add_formatted_text(p, text: str) -> None:
    """Render inline markdown (bold, italic, code, links) into a paragraph."""
    # Simple tokenizer: split by inline patterns and recurse on segments.
    pos = 0
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|(?<!\*)\*[^*]+\*(?!\*))"
    )
    for match in pattern.finditer(text):
        if match.start() > pos:
            _set_font(p.add_run(text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            _set_font(p.add_run(token[2:-2]), bold=True, color=HEADING_COLOR)
        elif token.startswith("`") and token.endswith("`"):
            run = p.add_run(token[1:-1])
            _set_font(run, name="Consolas", size=10, color=CODE_COLOR)
        elif token.startswith("["):
            m = LINK_RE.match(token)
            if m:
                run = p.add_run(m.group(1))
                _set_font(run, italic=True, color=HEADING_COLOR)
        elif token.startswith("*") and token.endswith("*"):
            _set_font(p.add_run(token[1:-1]), italic=True)
        else:
            _set_font(p.add_run(token))
        pos = match.end()
    if pos < len(text):
        _set_font(p.add_run(text[pos:]))


def _add_paragraph(doc: Document, text: str, italic: bool = False,
                   muted: bool = False, align_center: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if italic or muted:
        run = p.add_run(text)
        _set_font(run, italic=italic, color=MUTED_COLOR if muted else BODY_COLOR)
    else:
        _add_formatted_text(p, text)


def _add_list_item(doc: Document, text: str, ordered: bool = False,
                   level: int = 0) -> None:
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    _add_formatted_text(p, text)


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code)
    _set_font(run, name="Consolas", size=9, color=CODE_COLOR)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_font(run, size=10, bold=True, color=HEADING_COLOR)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            if c >= len(headers):
                continue
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_formatted_text(p, value)
            for run in p.runs:
                run.font.size = Pt(10)


def _parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    """Parse a markdown table starting at index `start`. Returns (headers, rows, next_index)."""
    def split_row(line: str) -> list[str]:
        parts = [c.strip() for c in line.strip().strip("|").split("|")]
        return parts

    headers = split_row(lines[start])
    # lines[start+1] is separator ---|---
    rows: list[list[str]] = []
    i = start + 2
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(split_row(lines[i]))
        i += 1
    return headers, rows, i


def md_to_docx(md_path: Path, out_path: Path, title: str) -> None:
    doc = Document()
    _configure_document(doc, title)

    # Cover title (first page, above the metadata block from the markdown).
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_font(run, size=24, bold=True, color=HEADING_COLOR)
    doc.add_paragraph()

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    i = 0
    in_code = False
    code_buffer: list[str] = []
    first_h1_seen = False

    while i < len(lines):
        line = lines[i]

        # Special HTML-comment markers
        stripped = line.strip()
        if stripped == "<!--PAGE_BREAK-->":
            doc.add_page_break()
            i += 1
            continue
        if stripped == "<!--WORD_TOC-->":
            _insert_toc_field(doc)
            doc.add_paragraph()
            i += 1
            continue

        # Skip raw HTML tags (portada uses <div align="center">, <br>, etc.)
        # These are rendered fine in HTML but would print as literal text in DOCX.
        if (stripped.startswith("<") and stripped.endswith(">")
                and not stripped.startswith("<!--")):
            i += 1
            continue

        # Fenced code blocks
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                _add_code_block(doc, "\n".join(code_buffer))
                in_code = False
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Skip the very first H1 (we already added a cover title)
        if line.startswith("# ") and not first_h1_seen:
            first_h1_seen = True
            i += 1
            continue

        # Headings
        if line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), 4)
            i += 1
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue
        if line.startswith("## "):
            _add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), 1)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        # Table
        if (line.strip().startswith("|") and i + 1 < len(lines)
                and re.match(r"^\|[\s\-:|]+\|\s*$", lines[i + 1])):
            headers, rows, next_i = _parse_table(lines, i)
            _add_table(doc, headers, rows)
            i = next_i
            continue

        # Unordered list
        if re.match(r"^\s*[-*]\s+", line):
            m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
            if m:
                indent_level = len(m.group(1)) // 2
                _add_list_item(doc, m.group(2), ordered=False, level=indent_level)
            i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
            if m:
                indent_level = len(m.group(1)) // 2
                _add_list_item(doc, m.group(2), ordered=True, level=indent_level)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            _add_paragraph(doc, line[2:], italic=True, muted=True)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Meta italic paragraph
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            _add_paragraph(doc, line.strip("*"), italic=True, muted=True,
                           align_center=True)
            i += 1
            continue

        # Regular paragraph
        _add_paragraph(doc, line)
        i += 1

    doc.save(out_path)
    print(f"DOCX generado: {out_path.name}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    for md_name, title in FILES:
        md_path = BASE_DIR / md_name
        stem = md_path.stem
        html_path = BASE_DIR / f"{stem}.html"
        docx_path = BASE_DIR / f"{stem}.docx"

        try:
            md_to_html(md_path, html_path, title)
            md_to_docx(md_path, docx_path, title)
        except PermissionError as e:
            print(f"AVISO: {stem}.docx está abierto en Word, se salta. ({e})")
            continue

    print("\nListo. Archivos en:", BASE_DIR)


if __name__ == "__main__":
    main()
